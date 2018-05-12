#! /usr/bin/env python3
from docx import Document
fileto_convert=open("raw_file.nct","r")
File=fileto_convert.read()
print(File)
fileto_convert.close()
fileto_convert=open("raw_file.nct","w")
document=Document(File)
for paragraph in document.paragraphs:
	for run in paragraph.runs:
			if run.font.name=="Preeti" and run.text!="":
				fileto_convert.write("%s%s"%(run.text,"\0"))
for table in document.tables:
	for row in table.rows: 
		for cell in row.cells:
   			for paragraph in cell.paragraphs:
   				for run in paragraph.runs:
   					if run.font.name=="Preeti"  and  run.text!="":
   						fileto_convert.write("%s%s"%(run.text,"\0"))
fileto_convert.close()
