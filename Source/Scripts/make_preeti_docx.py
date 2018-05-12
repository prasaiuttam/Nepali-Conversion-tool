#! /usr/bin/env python3
"""
	Author: Prashant Subedi
	This script is meant to be used as a part of file to file conversion function for unicode to preeti convertor.
"""
#for paragraphs
from docx import Document
import os
import re
converted_file=open("converted.nct","r")
raw_data=converted_file.read()
runs=[]
while(True):
	s=re.search("<#!new_run#!>",raw_data[0:len(raw_data)])
	if(s==None):
		break
	runs.append(raw_data[0:s.span()[0]])
	raw_data=raw_data[s.span()[1]:len(raw_data)]
Name=runs[0][6:len(runs[0])]
Folder=runs[1][13:len(runs[1])]
print(Folder)
i=2#Skiping the file and folder name
document=Document(Name)
for paragraph in document.paragraphs:
	for run in paragraph.runs:
		if len(run.text)>0:
			if(runs[i][len(runs[i])-1]=='1'):
				run.text=runs[i][0:len(runs[i])-1]
				run.font.name="Preeti"
			i+=1
for table in document.tables:
	for row in table.rows: 
		for cell in row.cells:
   			for paragraph in cell.paragraphs:
   				for run in paragraph.runs:
   					if len(run.text)>0:
	   					run.text=runs[i][0:len(runs[i])-1]
   						i+=1
   						if(runs[i][len(runs[i])-1]=='1'):
   							run.font.name="Preeti"
   						i+=1
print ("Document Conversion Complete")
print ("Cleaning Up")
os.remove("raw_file.nct")
os.remove("converted.nct")
Location=Folder+"/"+"converted.docx"
document.save(Location)

