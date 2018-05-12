#! /usr/bin/env python3
#for paragraphs
from docx import Document
import os
import re
converted_file=open("converted.nct","r")
raw_data=converted_file.read()
runs=[]
while(True):
	s=re.search("<#!new_run#!>",raw_data[0:len(raw_data)])
	if(s==None):
		break
	runs.append(raw_data[0:s.span()[0]])
	raw_data=raw_data[s.span()[1]:len(raw_data)]
Name=runs[0]
Folder=runs[1]
print(Folder)
i=2#Skiping the file and folder name
document=Document(Name)
for paragraph in document.paragraphs:
	for run in paragraph.runs:
		if(run.font.name=="Preeti"):
			run.text=runs[i]
			i+=1
			run.font.name=="Mangal"
for table in document.tables:
	for row in table.rows: 
		for cell in row.cells:
   			for paragraph in cell.paragraphs:
   				for run in paragraph.runs:
   					if(run.font.name=="Preeti"):
   						run.text=runs[i]
   						i+=1
   						run.font.name=="Mangal"
   									
print ("Document Conversion Complete")
print ("Cleaning Up")
os.remove("raw_file.nct")
os.remove("converted.nct")
Location=Folder+"/"+"converted.docx"
document.save(Location)

